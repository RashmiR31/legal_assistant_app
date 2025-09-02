"""
Legal Document Reviewer Assistant (single-file Streamlit app)

Features:
- Upload PDF / DOCX / TXT files via Streamlit
- Saves uploaded files to disk before ingestion
- Uses PyPDFLoader / UnstructuredWordDocumentLoader / python-docx fallback
- OCR fallback for scanned PDFs using pdf2image + pytesseract
- Splits text into chunks and indexes with OpenAI embeddings + FAISS
- Persists FAISS index to disk (safe only for local/trusted use)
- Retrieval QA + ConversationalRetrievalChain for audits
- Session-state persistence so buttons work across Streamlit reruns
- Helpful debug/log messages and error handling

Requirements (Python packages):
  pip install streamlit langchain openai faiss-cpu pypdf python-docx pytesseract pdf2image pillow
System packages:
  - tesseract (for OCR): e.g. `apt install tesseract-ocr`
  - poppler (for pdf2image): e.g. `apt install poppler-utils`

Usage:
  export OPENAI_API_KEY="sk-..."
  streamlit run legal_doc_reviewer_app.py

Notes:
- The FAISS loader requires `allow_dangerous_deserialization=True` when loading an index saved to disk. This is safe here only if you trust the index files (i.e., they were created by this app on your machine).
"""

import os
import shutil
import uuid
import traceback
from typing import List

import streamlit as st

# LangChain imports
from langchain.document_loaders import PyPDFLoader, UnstructuredWordDocumentLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from langchain_openai import ChatOpenAI
from langchain.chains import ConversationalRetrievalChain, RetrievalQA
from langchain.prompts import PromptTemplate

# Optional OCR imports - guard with try/except to provide helpful error messages
try:
    from pdf2image import convert_from_path
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

# Optional python-docx fallback
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# -----------------------------
# Config
# -----------------------------
UPLOAD_DIR = "uploads"
PERSIST_DIR = "faiss_index"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
EMBEDDING_MODEL = "text-embedding-3-small"  # controlled by OpenAIEmbeddings internally
LLM_MODEL = "gpt-3.5-turbo"  # change per your access & cost preferences

# -----------------------------
# Helpers: File handling & ingestion
# -----------------------------

def save_uploaded_files(uploaded_files) -> List[str]:
    """Save uploaded Streamlit files to disk and return list of paths."""
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    saved_paths = []
    for f in uploaded_files:
        safe_name = os.path.basename(f.name)
        dst = os.path.join(UPLOAD_DIR, safe_name)
        # overwrite if exists
        with open(dst, "wb") as out:
            shutil.copyfileobj(f, out)
        saved_paths.append(dst)
    return saved_paths


def ocr_pdf_to_docs(pdf_path: str, dpi=300) -> List[Document]:
    """Convert each page of PDF to image and OCR it into Documents.
    Returns a list of LangChain Document (one per page).
    """
    if not OCR_AVAILABLE:
        raise RuntimeError("OCR dependencies are not installed (pdf2image, pytesseract, pillow).")

    pages = convert_from_path(pdf_path, dpi=dpi)
    docs = []
    for i, page in enumerate(pages):
        text = pytesseract.image_to_string(page)
        metadata = {"source": pdf_path, "page": i + 1}
        docs.append(Document(page_content=text, metadata=metadata))
    return docs


def load_file_to_docs(file_path: str) -> List[Document]:
    """Load a file into a list of LangChain Document objects.
    Supports: PDF (text or scanned), DOCX (with fallbacks), TXT.
    """
    ext = os.path.splitext(file_path)[1].lower()
    docs: List[Document] = []

    if ext == ".pdf":
        # Attempt fast text extraction first
        try:
            loader = PyPDFLoader(file_path)
            docs = loader.load_and_split()
            # If loader returned very little text, assume scanned PDF and fallback to OCR
            combined = "".join([d.page_content.strip() for d in docs if d.page_content])
            if len(combined) < 40:
                if OCR_AVAILABLE:
                    st.info(f"PDF appears scanned or has little text; running OCR for {os.path.basename(file_path)}")
                    docs = ocr_pdf_to_docs(file_path)
                else:
                    st.warning(f"PDF {os.path.basename(file_path)} may be scanned. Install OCR deps (pdf2image, pytesseract).")
        except Exception as e:
            # Fallback: try OCR if available
            if OCR_AVAILABLE:
                st.info(f"PyPDFLoader failed for {file_path}, falling back to OCR: {e}")
                docs = ocr_pdf_to_docs(file_path)
            else:
                raise

    elif ext in [".docx", ".doc"]:
        # Try Unstructured loader, else python-docx fallback
        try:
            loader = UnstructuredWordDocumentLoader(file_path)
            docs = loader.load()
        except Exception:
            if DOCX_AVAILABLE:
                # simple fallback to python-docx
                docx = DocxDocument(file_path)
                text = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
                docs = [Document(page_content=text, metadata={"source": file_path})]
            else:
                raise

    elif ext in [".txt", ".md"]:
        loader = TextLoader(file_path, encoding="utf8")
        docs = loader.load()

    else:
        raise ValueError(f"Unsupported file type: {ext}")

    # Ensure metadata
    for d in docs:
        if "source" not in d.metadata:
            d.metadata["source"] = file_path
    return docs


def split_documents(docs: List[Document], chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(docs)


def build_faiss_index(docs: List[Document], persist_dir: str = PERSIST_DIR) -> FAISS:
    """Create or update a FAISS index from Documents. Persist to disk.
    Note: when loading from disk we use allow_dangerous_deserialization=True — only safe for local indexes you trust.
    """
    embeddings = OpenAIEmbeddings()
    if os.path.exists(persist_dir):
        try:
            index = FAISS.load_local(persist_dir, embeddings, allow_dangerous_deserialization=True)
            index.add_documents(docs)
            index.save_local(persist_dir)
        except Exception:
            # If load fails, create new index from scratch
            index = FAISS.from_documents(docs, embeddings)
            index.save_local(persist_dir)
    else:
        index = FAISS.from_documents(docs, embeddings)
        index.save_local(persist_dir)
    return index


def ingest_files(file_paths: List[str], persist_dir: str = PERSIST_DIR):
    """High-level pipeline: load files -> split -> index -> persist"""
    all_docs: List[Document] = []
    for path in file_paths:
        st.info(f"Loading {path} ...")
        docs = load_file_to_docs(path)
        st.info(f"Loaded {len(docs)} raw doc(s) from {os.path.basename(path)}")
        chunks = split_documents(docs)
        # augment metadata
        for c in chunks:
            if "chunk_id" not in c.metadata:
                c.metadata["chunk_id"] = str(uuid.uuid4())
        all_docs.extend(chunks)

    st.info(f"Total chunks to index: {len(all_docs)}")
    if not all_docs:
        st.warning("No text found in uploaded files. Check file types or install OCR dependencies for scanned PDFs.")
        return None

    index = build_faiss_index(all_docs, persist_dir=persist_dir)
    return index


# -----------------------------
# Streamlit app UI
# -----------------------------

st.set_page_config(page_title="Legal Document Reviewer Assistant", layout="wide")
st.title("⚖️ Legal Document Reviewer Assistant")

# Check OPENAI_API_KEY
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY") if hasattr(st, "secrets") else None
if not OPENAI_API_KEY:
    st.error("OPENAI_API_KEY not set. Please set the environment variable before running the app.")
    st.stop()

# Initialize session state
if "index" not in st.session_state:
    st.session_state.index = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "last_query" not in st.session_state:
    st.session_state.last_query = ""
if "last_answer" not in st.session_state:
    st.session_state.last_answer = ""
if "logs" not in st.session_state:
    st.session_state.logs = []

# Sidebar: upload + ingest
with st.sidebar.form(key="upload_form"):
    uploaded_files = st.file_uploader("Upload one or more documents", accept_multiple_files=True, type=["pdf", "docx", "txt"], help="Limit 200MB per file")
    ingest_btn = st.form_submit_button("Ingest / (re)build index")

if ingest_btn:
    if not uploaded_files:
        st.sidebar.warning("Please upload at least one document before ingesting.")
    else:
        try:
            # Save uploaded files and ingest
            saved = save_uploaded_files(uploaded_files)
            st.sidebar.info(f"Saved {len(saved)} files to '{UPLOAD_DIR}'")
            with st.spinner("Ingesting files and building index..."):
                idx = ingest_files(saved, persist_dir=PERSIST_DIR)
                if idx is None:
                    st.sidebar.error("Ingest produced no index (no text found). Check uploaded files or OCR dependencies.")
                else:
                    st.session_state.index = idx
                    st.sidebar.success("Ingest completed.")
        except Exception as e:
            st.sidebar.error(f"Ingest failed: {e}")
            st.sidebar.text(traceback.format_exc())

# Try to load index from disk if session doesn't have one
if st.session_state.index is None and os.path.exists(PERSIST_DIR):
    try:
        embeddings = OpenAIEmbeddings()
        st.session_state.index = FAISS.load_local(PERSIST_DIR, embeddings, allow_dangerous_deserialization=True)
        st.sidebar.success("Loaded existing index from disk.")
    except Exception as e:
        st.sidebar.error(f"Failed to load existing index: {e}")

index = st.session_state.index

if index is None:
    st.info("No index found. Upload documents and click 'Ingest / (re)build index' in the sidebar.")
    # show helpful troubleshooting info
    st.markdown("""
    **Troubleshooting**
    - If your PDFs are scanned images, install `tesseract-ocr` and `poppler-utils`.
    - Make sure `OPENAI_API_KEY` is set.
    """)
    st.stop()

# Build retriever + LLM chains
retriever = index.as_retriever(search_type="similarity", search_kwargs={"k": 4})
llm = ChatOpenAI(model=LLM_MODEL, temperature=0)

simple_qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever, return_source_documents=True)

qa_chain = ConversationalRetrievalChain.from_llm(
    llm=llm,
    retriever=retriever,
    condense_question_prompt=PromptTemplate.from_template(
        "Given the chat history and a follow-up question, rewrite it to be standalone:\n\n{chat_history}\n\nFollow-up: {question}\n\nStandalone:"
    ),
    return_source_documents=True,
)

# Main UI: Ask
st.header("Ask a question about the uploaded documents")
question = st.text_input("Type your question here (e.g. 'Is there a termination clause? Summarize it.')")

col1, col2 = st.columns([3, 1])
with col2:
    ask_btn = st.button("Ask")

if ask_btn:
    if not question or question.strip() == "":
        st.warning("Please enter a question.")
    else:
        st.session_state.last_query = question
        try:
            with st.spinner("Retrieving answer from documents..."):
                try:
                    result = simple_qa({"query": question})
                    answer = result.get("result") or result.get("answer") or str(result)
                    source_docs = result.get("source_documents", [])
                except Exception as e:
                    st.error(f"Failed to get answer: {e}")
                    st.text(traceback.format_exc())

            st.subheader("Answer")
            st.write(answer)
            st.session_state.last_answer = answer

            st.subheader("Source snippets")
            if not source_docs:
                st.info("No source snippets returned. Try a different question or increase 'k' in retriever settings.")
            for doc in source_docs:
                src = doc.metadata.get("source", "unknown")
                page = doc.metadata.get("page")
                chunk_id = doc.metadata.get("chunk_id")
                header = f"Source: {os.path.basename(src)}"
                if page:
                    header += f" | page: {page}"
                if chunk_id:
                    header += f" | chunk: {chunk_id}"
                with st.expander(header):
                    st.write(doc.page_content)

            # Save to chat history
            st.session_state.chat_history.append((question, answer))

        except Exception as e:
            st.error(f"Failed to get answer: {e}")
            st.text(traceback.format_exc())

# Audit feature
st.header("Audit / Suggest missing clauses and structure")
if st.button("Run Document Audit"):
    with st.spinner("Running audit (quick scan + LLM)..."):
        try:
            # Quick keyword-based scan
            keywords = {
                "Termination": ["termination", "end of agreement", "terminate"],
                "Governing Law": ["governing law", "jurisdiction"],
                "Confidentiality": ["confidential", "nondisclosure", "non-disclosure", "NDA"],
                "Indemnity": ["indemnify", "hold harmless", "indemnity"],
                "Limitation of Liability": ["limitation of liability", "liability cap", "cap on liability"],
                "Definitions": ["definitions", "meaning of"],
                "Force Majeure": ["force majeure", "act of god"],
                "Payment Terms": ["payment", "invoice", "due date", "fees"],
                "Assignment": ["assignment", "assign", "successors"],
            }

            present = []
            missing = []
            # For each clause, search and look for keywords in retrieved docs
            for name, kws in keywords.items():
                q = f"Search the documents for the clause about {name}. Return any short snippet if present." 
                res = simple_qa({"query": q})
                # Res might be a dict or a string; try to get text
                text = ""
                if isinstance(res, dict):
                    text = (res.get("result") or res.get("answer") or "").lower()
                else:
                    text = str(res).lower()

                if any(k.lower() in text for k in kws):
                    present.append(name)
                else:
                    missing.append(name)

            st.subheader("Quick keyword scan")
            st.write("Present:", present)
            st.write("Missing (quick scan):", missing)

            # LLM-based audit: ask the chain to produce a checklist + drafts
            audit_prompt = f"You are an expert legal reviewer. Audit the uploaded documents and: 1) produce a checklist of common commercial contract clauses and mark present/missing, 2) for missing items provide a short rationale and a suggested clause draft (3-6 sentences).\n".replace("\\n", "\n")

            audit_response = qa_chain({
            "question": "Perform a contract audit and suggest missing clauses and drafts.",
            "chat_history": []
            })
            audit_result = audit_response.get("answer") or audit_response.get("result") or str(audit_response)

            st.subheader("LLM Audit (detailed)")
            st.write(audit_result)

            # Export audit text as download
            audit_text = "Quick scan present: " + ", ".join(present) + "\nQuick scan missing: " + ", ".join(missing) + "\n\nLLM Audit:\n" + str(audit_result)
            st.download_button("Download audit report (TXT)", data=audit_text, file_name="audit_report.txt")

        except Exception as e:
            st.error(f"Audit failed: {e}")
            st.text(traceback.format_exc())

# Sidebar actions and info
with st.sidebar:
    st.write("### Export & Settings")
    if st.button("Export index snapshot"):
        if os.path.exists(PERSIST_DIR):
            st.success(f"Index snapshot saved in `{PERSIST_DIR}`")
        else:
            st.warning("No index found on disk to export.")

    if st.button("Clear index & uploads (danger)"):
        try:
            if os.path.exists(PERSIST_DIR):
                shutil.rmtree(PERSIST_DIR)
            if os.path.exists(UPLOAD_DIR):
                shutil.rmtree(UPLOAD_DIR)
            st.session_state.index = None
            st.success("Cleared index and uploads. Reload the app to start fresh.")
        except Exception as e:
            st.error(f"Failed to clear: {e}")

    st.markdown("""
    **Notes**
    - Uploaded documents are processed with OpenAI embeddings. Ensure `OPENAI_API_KEY` is set.
    - For scanned PDFs, install `tesseract` and `poppler-utils` on the server.
    - Only load indexes you created yourself (we use `allow_dangerous_deserialization=True`).
    """)

# End of app
